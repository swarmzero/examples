import msal
import os
import logging
import requests
from urllib.parse import quote_plus
from dotenv import load_dotenv
from docx import Document
from io import BytesIO


logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

load_dotenv()


def publish_to_sharepoint(title: str, results_text: str) -> str:
    """
    Publishes results to SharePoint by creating a new file in the specified drive.

    Args:
        title (str): The title of the file to be uploaded.
        results_text (str): The content to be uploaded as a file.

    Returns:
        str: The URL of the uploaded SharePoint file.

    Raises:
        Exception: If an error occurs during authentication or file upload.
    """

    SHAREPOINT_CLIENT_ID = os.getenv("SHAREPOINT_CLIENT_ID")
    SHAREPOINT_CLIENT_SECRET = os.getenv("SHAREPOINT_CLIENT_SECRET")
    SHAREPOINT_TENANT_ID = os.getenv("SHAREPOINT_TENANT_ID")
    SHAREPOINT_SITE_ID = os.getenv("SHAREPOINT_SITE_ID")
    SHAREPOINT_DRIVE_ID = os.getenv("SHAREPOINT_DRIVE_ID")

    try:
        logger.debug("Initiating SharePoint authentication")
        authority = f"https://login.microsoftonline.com/{SHAREPOINT_TENANT_ID}"
        logger.debug(f"Using authority URL: {authority}")

        app = msal.ConfidentialClientApplication(
            SHAREPOINT_CLIENT_ID,
            authority=authority,
            client_credential=SHAREPOINT_CLIENT_SECRET,
        )

        scopes = ["https://graph.microsoft.com/.default"]
        logger.debug(f"Requesting scopes: {scopes}")

        # get access token
        logger.info("Requesting access token")
        result = app.acquire_token_for_client(scopes=scopes)

        if "access_token" not in result:
            error_msg = result.get("error_description", "Unknown error")
            error_code = result.get("error", "Unknown error code")
            logger.error(
                f"""Failed to obtain SharePoint access token:
                Error Code: {error_code}
                Description: {error_msg}
                Full Result: {result}"""
            )
            return None

        logger.info("Successfully acquired access token")
        access_token = result["access_token"]
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        logger.debug("Request headers prepared (token hidden)")

        file_name = f"{title}.docx"

        # convert text content to basic Word document format
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(results_text)

        # save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        file_content = doc_bytes.getvalue()

        logger.info(f"Preparing file upload: {file_name}")
        logger.debug(f"File size: {len(file_content)} bytes")

        upload_url = (
            f"https://graph.microsoft.com/v1.0/sites/{SHAREPOINT_SITE_ID}"
            f"/drives/{SHAREPOINT_DRIVE_ID}/root:/{quote_plus(file_name)}:/content"
        )
        logger.debug(f"Upload URL: {upload_url}")

        logger.info(f"Uploading file to SharePoint: {file_name}")
        response = requests.put(upload_url, headers=headers, data=file_content)

        logger.debug(
            f"""SharePoint Response Details:
            Status Code: {response.status_code}
            Headers: {dict(response.headers)}
            Response Text: {response.text}"""
        )

        if response.status_code in [200, 201]:
            try:
                file_info = response.json()
                web_url = file_info.get("webUrl")
                logger.info(f"File uploaded successfully")
                logger.info(f"Web URL: {web_url}")
                return web_url
            except Exception as e:
                logger.error(
                    f"""Failed to parse successful upload response:
                    Error: {str(e)}
                    Response Text: {response.text}"""
                )
                return None
        else:
            logger.error(
                f"""SharePoint upload failed:
                Status Code: {response.status_code}
                Response: {response.text}
                URL: {upload_url}
                File Name: {file_name}"""
            )
            return None

    except Exception as e:
        logger.error("Unexpected error during SharePoint upload", exc_info=True)
        logger.debug(f"Error details: {str(e)}")
        return None
